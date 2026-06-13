import asyncio
import io
import re

import pdfplumber
from docx import Document

from backend.schemas.resume_parser import ResumeParsedData, ResumeParseResponse

# 配置
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {".pdf", ".docx"}
TIMEOUT_SECONDS = 30


def validate_file(filename: str, file_size: int) -> str | None:
    """前置校验：扩展名 + 文件大小，返回错误信息或 None"""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return f"不支持的文件格式「{ext}」，仅支持 PDF 和 DOCX。如需上传其他格式，请先转换为 PDF 或 DOCX"
    if file_size > MAX_FILE_SIZE:
        return f"文件过大（{file_size // 1024 // 1024}MB），最大支持 10MB"
    return None


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """从 PDF 提取文本"""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 DOCX 提取文本"""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    return "\n".join(text_parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """根据文件类型提取文本"""
    ext = "." + filename.rsplit(".", 1)[-1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    return ""


def _parse_by_regex(text: str) -> dict | None:
    """正则解析简历文本为结构化数据"""
    result = {
        "basic_info": {},
        "education": [],
        "internship_exp": [],
        "project_exp": [],
        "personal_strengths": [],
    }

    lines = text.split("\n")
    lines = [line.strip() for line in lines if line.strip()]

    # ========== 基本信息 ==========
    # 姓名（通常是第一行或包含"姓名"的行）
    for line in lines[:5]:
        m = re.search(r"姓名[：:]\s*(.+)", line)
        if m:
            result["basic_info"]["name"] = m.group(1).strip()
            break
    if not result["basic_info"].get("name") and lines:
        # 第一行作为姓名（如果不太长）
        first = lines[0].strip()
        if len(first) <= 10 and not any(kw in first for kw in ["简历", "求职", "个人"]):
            result["basic_info"]["name"] = first

    # 电话
    phone_match = re.search(r"(?:电话|手机|Tel)[：:]?\s*([\d\-\sxX]+)", text)
    if phone_match:
        result["basic_info"]["phone"] = phone_match.group(1).strip()

    # 邮箱
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    if email_match:
        result["basic_info"]["email"] = email_match.group(0).strip()

    # ========== 教育经历 ==========
    edu_section = _extract_section(text, ["教育经历", "教育背景", "学历"])
    if edu_section:
        edu_items = _parse_education(edu_section)
        result["education"] = edu_items

    # ========== 实习经历 ==========
    intern_section = _extract_section(text, ["实习经历", "工作经历", "工作经验", "实习"])
    if intern_section:
        intern_items = _parse_experience(intern_section, "internship")
        result["internship_exp"] = intern_items

    # ========== 项目经历 ==========
    proj_section = _extract_section(text, ["项目经历", "项目经验", "项目"])
    if proj_section:
        proj_items = _parse_experience(proj_section, "project")
        result["project_exp"] = proj_items

    # ========== 个人优势 ==========
    strength_section = _extract_section(text, ["个人优势", "个人技能", "技能特长", "自我评价", "专业技能", "个人总结"])
    if strength_section:
        items = _parse_list_items(strength_section)
        result["personal_strengths"] = items

    # 检查是否解析到了有效内容
    has_content = (
        result["basic_info"]
        or result["education"]
        or result["internship_exp"]
        or result["project_exp"]
        or result["personal_strengths"]
    )
    return result if has_content else None


# 所有可能的章节标题（用于分隔）
ALL_SECTIONS = [
    "教育经历", "教育背景", "学历",
    "实习经历", "工作经历", "工作经验", "实习",
    "项目经历", "项目经验", "项目",
    "个人优势", "个人技能", "技能特长", "自我评价", "专业技能", "个人总结",
]


def _extract_section(text: str, headers: list[str]) -> str:
    """提取指定章节的内容，用已知标题列表做分隔"""
    for header in headers:
        idx = text.find(header)
        if idx == -1:
            continue
        start = idx + len(header)
        while start < len(text) and text[start] in "：: \t":
            start += 1
        remaining = text[start:]

        # 找下一个已知章节标题的位置
        earliest = len(remaining)
        for other_header in ALL_SECTIONS:
            if other_header in headers:
                continue  # 跳过当前章节自己的标题
            other_idx = remaining.find(other_header)
            if other_idx != -1 and other_idx < earliest:
                earliest = other_idx

        return remaining[:earliest].strip()
    return ""


def _parse_education(section: str) -> list[dict]:
    """解析教育经历"""
    items = []
    # 按行分析，寻找学校相关关键词
    lines = section.split("\n")
    current = {}
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 匹配学校行（包含"大学"、"学院"、"学校"等）
        if re.search(r"(?:大学|学院|学校|高校|本科|硕士|博士)", line):
            if current:
                items.append(current)
            current = {"school": "", "degree": "", "major": "", "time": "", "courses": ""}
            # 提取学校名和学位/专业
            parts = re.split(r"[·|｜,，]", line)
            if parts:
                current["school"] = parts[0].strip()
            # 提取学位
            degree_match = re.search(r"(本科|硕士|博士|专科|学士|MBA)", line)
            if degree_match:
                current["degree"] = degree_match.group(1)
            # 提取专业
            major_match = re.search(r"专业[：:]\s*(.+?)(?:\s|$)", line)
            if major_match:
                current["major"] = major_match.group(1).strip()
            # 提取时间
            time_match = re.search(r"(\d{4}[\.\-/]\d{1,2}[\s\-~—至]+\d{4}[\.\-/]\d{1,2}|\d{4}[\.\-/]\d{1,2}\s*[-~]\s*(?:至今|现在|present))", line, re.IGNORECASE)
            if time_match:
                current["time"] = time_match.group(1).strip()
        elif current:
            # 补充信息（课程等）
            if "课程" in line or "主修" in line:
                current["courses"] = line.split("：")[-1].strip() if "：" in line else line.split(":")[-1].strip()
            elif not current.get("time"):
                time_match = re.search(r"(\d{4}[\.\-/]\d{1,2}[\s\-~—至]+\d{4}[\.\-/]\d{1,2}|\d{4}[\.\-/]\d{1,2}\s*[-~]\s*(?:至今|现在))", line, re.IGNORECASE)
                if time_match:
                    current["time"] = time_match.group(1).strip()
    if current:
        items.append(current)
    return items


def _parse_experience(section: str, exp_type: str) -> list[dict]:
    """解析实习/项目经历"""
    items = []
    lines = section.split("\n")
    current = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 检测是否是新条目（包含 | 或时间格式）
        is_new = False
        if "|" in line or "｜" in line:
            is_new = True
        elif re.search(r"\d{4}[\.\-/]\d{1,2}", line) and len(line) < 80:
            is_new = True

        if is_new:
            if current:
                items.append(current)
            current = {"name": "", "role": "", "company": "", "time": "", "description": []}
            # 提取时间
            time_match = re.search(r"(\d{4}[\.\-/]\d{1,2}[\s\-~—至]+\d{4}[\.\-/]\d{1,2}|\d{4}[\.\-/]\d{1,2}\s*[-~]\s*(?:至今|现在|present))", line, re.IGNORECASE)
            if time_match:
                current["time"] = time_match.group(1).strip()
            # 去掉时间部分后再拆分
            clean_line = line[:time_match.start()].strip() if time_match else line
            parts = re.split(r"[|｜]", clean_line)
            if exp_type == "internship":
                if len(parts) >= 1:
                    current["company"] = parts[0].strip()
                if len(parts) >= 2:
                    current["role"] = parts[1].strip()
            else:
                if len(parts) >= 1:
                    current["name"] = parts[0].strip()
                if len(parts) >= 2:
                    current["role"] = parts[1].strip()
        elif current:
            # 描述行
            desc_line = re.sub(r"^[\s]*[-•·*\d.、)）]+\s*", "", line)
            if desc_line:
                current["description"].append(desc_line)

    if current:
        items.append(current)
    return items


def _parse_list_items(section: str) -> list[str]:
    """解析列表项"""
    items = []
    lines = section.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 去掉列表标记
        cleaned = re.sub(r"^[\s]*[-•·*\d.、)）]+\s*", "", line)
        if cleaned:
            items.append(cleaned)
    return items


async def _llm_parse(text: str) -> dict | None:
    """LLM 兜底解析（模拟，后续接入真实 LLM）"""
    # TODO: 接入真实 LLM API
    # 当前返回简单的行分类结果
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return None

    result = {
        "basic_info": {"name": lines[0] if len(lines[0]) <= 10 else ""},
        "education": [],
        "internship_exp": [],
        "project_exp": [],
        "personal_strengths": [],
    }

    # 简单按行归类
    current_section = "personal_strengths"
    for line in lines[1:]:
        if any(kw in line for kw in ["教育", "学校", "大学"]):
            current_section = "education"
        elif any(kw in line for kw in ["实习", "工作", "公司"]):
            current_section = "internship_exp"
        elif any(kw in line for kw in ["项目", "开发"]):
            current_section = "project_exp"
        elif any(kw in line for kw in ["技能", "优势", "自我"]):
            current_section = "personal_strengths"
        elif len(line) > 3:
            if current_section == "personal_strengths":
                result["personal_strengths"].append(line)

    return result


async def parse_resume(filename: str, file_bytes: bytes) -> ResumeParseResponse:
    """解析简历文件，带超时控制"""
    # 前置校验
    validation_error = validate_file(filename, len(file_bytes))
    if validation_error:
        return ResumeParseResponse(success=False, error=validation_error)

    try:
        # 提取文本（带超时）
        try:
            raw_text = await asyncio.wait_for(
                asyncio.to_thread(extract_text, filename, file_bytes),
                timeout=TIMEOUT_SECONDS,
            )
        except Exception:
            return ResumeParseResponse(
                success=False,
                error="文件无法读取，请检查是否加密或损坏",
            )

        # 检查内容是否为空
        if not raw_text or len(raw_text.strip()) < 10:
            return ResumeParseResponse(
                success=False,
                error="未识别到有效内容，请检查文件是否可编辑",
            )

        # 正则解析
        result = await asyncio.wait_for(
            asyncio.to_thread(_parse_by_regex, raw_text),
            timeout=TIMEOUT_SECONDS,
        )

        if result:
            return ResumeParseResponse(
                success=True,
                data=ResumeParsedData(**result),
                method="regex",
                raw_text=raw_text[:500],  # 只返回前500字符用于调试
            )

        # LLM 兜底
        result = await asyncio.wait_for(
            _llm_parse(raw_text),
            timeout=TIMEOUT_SECONDS,
        )

        if result:
            return ResumeParseResponse(
                success=True,
                data=ResumeParsedData(**result),
                method="llm",
                raw_text=raw_text[:500],
            )

        return ResumeParseResponse(
            success=False,
            error="无法解析简历内容，请检查文件格式",
            raw_text=raw_text[:500],
        )

    except asyncio.TimeoutError:
        return ResumeParseResponse(
            success=False,
            error=f"解析超时（{TIMEOUT_SECONDS}秒），请稍后重试",
        )
    except Exception as e:
        return ResumeParseResponse(
            success=False,
            error=f"解析异常：{str(e)}",
        )
