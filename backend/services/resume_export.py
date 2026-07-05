"""简历导出服务 — PDF / DOCX 生成"""

import io
import os
import platform

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------- 中文字体注册 ----------

_FONT_REGISTERED = False
_FONT_NAME = "SimHei"


def _register_chinese_font():
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return

    candidates = []
    if platform.system() == "Windows":
        font_dir = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
        candidates = [
            os.path.join(font_dir, "simhei.ttf"),   # 黑体
            os.path.join(font_dir, "simsun.ttc"),    # 宋体
            os.path.join(font_dir, "msyh.ttc"),      # 微软雅黑
        ]
    elif platform.system() == "Darwin":
        candidates = [
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/PingFang.ttc",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        ]

    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(_FONT_NAME, path))
                _FONT_REGISTERED = True
                return
            except Exception:
                continue

    # 回退：尝试用系统默认字体（中文会乱码但不会报错）
    _FONT_REGISTERED = True


# ---------- PDF 样式 ----------

def _pdf_styles():
    return {
        "name": ParagraphStyle(
            "Name", fontName=_FONT_NAME, fontSize=18, alignment=TA_CENTER,
            spaceAfter=2 * mm, leading=24,
        ),
        "contact": ParagraphStyle(
            "Contact", fontName=_FONT_NAME, fontSize=9, alignment=TA_CENTER,
            textColor="#666666", spaceAfter=4 * mm, leading=14,
        ),
        "section_title": ParagraphStyle(
            "SectionTitle", fontName=_FONT_NAME, fontSize=12,
            spaceBefore=4 * mm, spaceAfter=2 * mm, leading=16,
        ),
        "item_title": ParagraphStyle(
            "ItemTitle", fontName=_FONT_NAME, fontSize=10,
            spaceBefore=2 * mm, spaceAfter=1 * mm, leading=14,
        ),
        "item_subtitle": ParagraphStyle(
            "ItemSubtitle", fontName=_FONT_NAME, fontSize=9,
            textColor="#666666", spaceAfter=1 * mm, leading=13,
        ),
        "body": ParagraphStyle(
            "Body", fontName=_FONT_NAME, fontSize=9,
            textColor="#444444", leading=14,
        ),
        "bullet": ParagraphStyle(
            "Bullet", fontName=_FONT_NAME, fontSize=9,
            textColor="#444444", leftIndent=12, leading=14,
            bulletIndent=0, spaceBefore=1 * mm,
        ),
    }


# ---------- PDF 生成 ----------


def generate_pdf(resume_data: dict) -> bytes:
    """将简历数据渲染为 PDF，返回 bytes"""
    _register_chinese_font()
    styles = _pdf_styles()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )

    story = []
    basic = resume_data.get("basic_info") or {}

    # 姓名
    name = basic.get("name", "")
    if name:
        story.append(Paragraph(_escape(name), styles["name"]))

    # 联系方式
    contacts = []
    if basic.get("email"):
        contacts.append(basic["email"])
    if basic.get("phone"):
        contacts.append(basic["phone"])
    if contacts:
        story.append(Paragraph(" | ".join(contacts), styles["contact"]))

    story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc", spaceAfter=3 * mm))

    # 教育经历
    edu_list = resume_data.get("education") or []
    if edu_list:
        story.append(Paragraph(_escape("教育经历"), styles["section_title"]))
        story.append(HRFlowable(width="100%", thickness=0.3, color="#e0e0e0", spaceAfter=2 * mm))
        for edu in edu_list:
            title = edu.get("school", "")
            degree = edu.get("degree", "")
            major = edu.get("major", "")
            time_range = edu.get("time", "")
            left = f"{_escape(title)}  {_escape(degree)}  {_escape(major)}"
            story.append(Paragraph(left, styles["item_title"]))
            if time_range:
                story.append(Paragraph(_escape(time_range), styles["item_subtitle"]))
            courses = edu.get("courses", "")
            if courses:
                story.append(Paragraph(f"主修课程：{_escape(courses)}", styles["body"]))

    # 实习经历
    intern_list = resume_data.get("internship_exp") or []
    if intern_list:
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(_escape("实习经历"), styles["section_title"]))
        story.append(HRFlowable(width="100%", thickness=0.3, color="#e0e0e0", spaceAfter=2 * mm))
        for intern in intern_list:
            company = intern.get("company", "")
            role = intern.get("role", "")
            time_range = intern.get("time", "")
            story.append(Paragraph(f"{_escape(company)} — {_escape(role)}", styles["item_title"]))
            if time_range:
                story.append(Paragraph(_escape(time_range), styles["item_subtitle"]))
            for desc in intern.get("description") or []:
                story.append(Paragraph(f"• {_escape(desc)}", styles["bullet"]))

    # 项目经历
    proj_list = resume_data.get("project_exp") or []
    if proj_list:
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(_escape("项目经历"), styles["section_title"]))
        story.append(HRFlowable(width="100%", thickness=0.3, color="#e0e0e0", spaceAfter=2 * mm))
        for proj in proj_list:
            proj_name = proj.get("name", "")
            role = proj.get("role", "")
            time_range = proj.get("time", "")
            story.append(Paragraph(f"{_escape(proj_name)} — {_escape(role)}", styles["item_title"]))
            if time_range:
                story.append(Paragraph(_escape(time_range), styles["item_subtitle"]))
            for desc in proj.get("description") or []:
                story.append(Paragraph(f"• {_escape(desc)}", styles["bullet"]))

    # 个人优势
    strengths = resume_data.get("personal_strengths") or []
    if strengths:
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(_escape("个人优势"), styles["section_title"]))
        story.append(HRFlowable(width="100%", thickness=0.3, color="#e0e0e0", spaceAfter=2 * mm))
        for s in strengths:
            story.append(Paragraph(f"• {_escape(s)}", styles["bullet"]))

    doc.build(story)
    return buf.getvalue()


# ---------- DOCX 生成 ----------

FONT_NAME = "宋体"


def _set_run_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None):
    """统一设置 run 的字体（西文 + 东亚）"""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_bottom_border(paragraph, color: str = "999999", size: int = 4):
    """给段落添加底部边框作为分割线（自适应页面宽度）"""
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(doc, text: str, size: float = 10, bold: bool = False, color: RGBColor | None = None):
    """添加段落并返回 run（统一字体）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    return p, run


def _add_item_header(doc, title: str, time_range: str = ""):
    """添加条目标题行：标题左对齐加粗，时间右对齐同行（双列表格）"""
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # 去除表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)

    # 左列：标题
    cell_left = table.cell(0, 0)
    cell_left.width = Cm(12)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.space_before = Pt(2)
    p_left.paragraph_format.space_after = Pt(1)
    run_title = p_left.add_run(title)
    _set_run_font(run_title, size=11, bold=True)

    # 右列：时间
    cell_right = table.cell(0, 1)
    cell_right.width = Cm(3.92)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(2)
    p_right.paragraph_format.space_after = Pt(1)
    run_time = p_right.add_run(time_range or "")
    _set_run_font(run_time, size=9, color=RGBColor(0x99, 0x99, 0x99))


def generate_docx(resume_data: dict) -> bytes:
    """将简历数据渲染为 DOCX，返回 bytes"""
    doc = Document()

    # A4 纸张设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # 默认样式：宋体
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    basic = resume_data.get("basic_info") or {}

    # 姓名
    name = basic.get("name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        _set_run_font(run, size=22, bold=True)

    # 联系方式
    contacts = []
    if basic.get("email"):
        contacts.append(basic["email"])
    if basic.get("phone"):
        contacts.append(basic["phone"])
    if contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contacts))
        _set_run_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))

    # 分割线（段落底部边框，自适应宽度）
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(4)
    hr.paragraph_format.space_after = Pt(4)
    _add_bottom_border(hr)

    # 教育经历
    edu_list = resume_data.get("education") or []
    if edu_list:
        _add_section_header(doc, "教育经历")
        for edu in edu_list:
            school = edu.get("school", "")
            degree = edu.get("degree", "")
            major = edu.get("major", "")
            _add_item_header(doc, f"{school}  {degree}  {major}", edu.get("time", ""))
            courses = edu.get("courses", "")
            if courses:
                _add_run(doc, f"主修课程：{courses}")

    # 实习经历
    intern_list = resume_data.get("internship_exp") or []
    if intern_list:
        _add_section_header(doc, "实习经历")
        for intern in intern_list:
            company = intern.get("company", "")
            role = intern.get("role", "")
            _add_item_header(doc, f"{company} — {role}", intern.get("time", ""))
            for desc in intern.get("description") or []:
                _add_run(doc, f"• {desc}")

    # 项目经历
    proj_list = resume_data.get("project_exp") or []
    if proj_list:
        _add_section_header(doc, "项目经历")
        for proj in proj_list:
            proj_name = proj.get("name", "")
            role = proj.get("role", "")
            _add_item_header(doc, f"{proj_name} — {role}", proj.get("time", ""))
            for desc in proj.get("description") or []:
                _add_run(doc, f"• {desc}")

    # 个人优势
    strengths = resume_data.get("personal_strengths") or []
    if strengths:
        _add_section_header(doc, "个人优势")
        for s in strengths:
            _add_run(doc, f"• {s}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section_header(doc: Document, title: str):
    """添加段落标题 + 底部分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    _set_run_font(run, size=13, bold=True)
    _add_bottom_border(p, color="cccccc", size=4)


def _escape(text: str) -> str:
    """转义 XML 特殊字符（reportlab Paragraph 需要）"""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))
